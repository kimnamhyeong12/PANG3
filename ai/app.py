"""
사하구 현장 보고서 AI 엔진
- 입력: JSON (파일 또는 sys.argv[1])
- 출력: HWP 호환 docx + stdout 마지막 줄 @@AI_RESULT@@ JSON
"""

import json
import os
import re
import sys
from datetime import date
from pathlib import Path

from dotenv import load_dotenv
from google import genai
from PIL import Image, ImageOps
from pillow_heif import register_heif_opener
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


register_heif_opener()

AI_DIR = Path(__file__).resolve().parent
load_dotenv(AI_DIR / ".env")

API_KEY = os.environ.get("GEMINI_API_KEY")
client = genai.Client(api_key=API_KEY) if API_KEY else None

RESULT_PREFIX = "@@AI_RESULT@@"



import urllib.request
from urllib.parse import urlparse

def download_if_url(source_path: str | None, download_dir: Path) -> str | None:
    """
    들어온 경로가 S3 웹 주소(http/https) 형태이면 파일을 다운로드하고 로컬 경로를 반환합니다.
    원래 로컬 경로이거나 없으면 그대로 통과시킵니다.
    """
    if not source_path:
        return None
        
    # http:// 나 https:// 로 시작하는 웹 주소인지 검사
    parsed_url = urlparse(source_path)
    if parsed_url.scheme in ('http', 'https'):
        try:
            download_dir.mkdir(parents=True, exist_ok=True)
            # URL 주소 맨 뒤에서 실제 파일명(예: photo_0.jpg) 추출
            file_name = os.path.basename(parsed_url.path) or "downloaded_img.jpg"
            local_target_path = download_dir / file_name
            
            print(f"[S3 연동] 웹 이미지 감지됨. 다운로드 시작: {source_path}", file=sys.stderr)
            urllib.request.urlretrieve(source_path, str(local_target_path))
            print(f"[S3 연동] 다운로드 완료 -> {local_target_path}", file=sys.stderr)
            
            return str(local_target_path.resolve())
        except Exception as e:
            print(f"[S3 다운로드 에러] {source_path}: {e}", file=sys.stderr)
            return None
            
    return source_path


# ---------------------------------------------------------------------------
# 경로 · 입력 로드
# ---------------------------------------------------------------------------

def resolve_path(path: str | None) -> str | None:
    if not path:
        return None
    p = Path(path)
    if not p.is_absolute():
        p = AI_DIR / p
    return str(p.resolve()) if p.exists() else None


def load_payload() -> dict:
    if len(sys.argv) <= 1:
        sample = AI_DIR / "samples" / "garak_sample.json"
        return json.loads(sample.read_text(encoding="utf-8"))
    arg = sys.argv[1]
    if arg.endswith(".json"):
        json_path = Path(arg)
        if not json_path.is_absolute():
            json_path = AI_DIR / json_path
        if json_path.is_file():
            return json.loads(json_path.read_text(encoding="utf-8"))
    return json.loads(arg)


# ---------------------------------------------------------------------------
# 이미지 전처리
# ---------------------------------------------------------------------------

def sanitize_image(
    source_path: str | None,
    target_path: str,
    canvas_size: tuple[int, int] = (1800, 1200),
    fit_mode: str = "cover",
) -> str | None:
    resolved = resolve_path(source_path) if source_path and not os.path.isabs(source_path) else source_path
    if resolved and not os.path.isabs(resolved):
        resolved = resolve_path(resolved)
    if not resolved or not os.path.exists(resolved):
        return None
    try:
        with Image.open(resolved) as img:
            img = ImageOps.exif_transpose(img).convert("RGB")
            if fit_mode == "contain":
                canvas = Image.new("RGB", canvas_size, (255, 255, 255))
                img.thumbnail(canvas_size, Image.Resampling.LANCZOS)
                left = (canvas_size[0] - img.width) // 2
                top = (canvas_size[1] - img.height) // 2
                canvas.paste(img, (left, top))
            else:
                canvas = ImageOps.fit(
                    img,
                    canvas_size,
                    method=Image.Resampling.LANCZOS,
                    centering=(0.5, 0.5),
                )
            canvas.save(target_path, "JPEG", quality=92)
        return target_path
    except Exception as e:
        print(f"[Image Sanitizer Error] {source_path}: {e}", file=sys.stderr)
        return None


# ---------------------------------------------------------------------------
# 사진 코멘트 파싱 · 섹션 그룹
# ---------------------------------------------------------------------------

def normalize_section(section: str) -> str:
    section = section.strip()
    m = re.match(r"^현장사진\s*\(\s*(\d+)\s*\)$", section, re.IGNORECASE)
    if m:
        return f"현장사진 ({m.group(1)})"
    m = re.match(r"^현장사진\s*\(\s*(\d+)\s*\)", section, re.IGNORECASE)
    if m:
        return f"현장사진 ({m.group(1)})"
    if section in ("전", "중", "후"):
        return f"현장사진 ({section})"
    return section


def parse_photo_comment(text: str, auto_index: int) -> tuple[str, str]:
    text = (text or "").strip()
    if not text:
        return f"__auto_{auto_index}__", ""

    if "|" in text:
        section, caption = text.split("|", 1)
        return normalize_section(section.strip()), caption.strip()

    lines = [ln.strip() for ln in text.splitlines() if ln.strip()]
    if len(lines) >= 2:
        return normalize_section(lines[0]), "\n".join(lines[1:])

    if text in ("전", "중", "후"):
        return f"현장사진 ({text})", ""

    m = re.match(r"^현장사진\s*\(\s*(\d+)\s*\)\s*$", text, re.IGNORECASE)
    if m:
        return f"현장사진 ({m.group(1)})", ""

    m = re.match(r"^현장사진\s*\(\s*(\d+)\s*\)\s*(.+)$", text, re.IGNORECASE | re.DOTALL)
    if m:
        return f"현장사진 ({m.group(1)})", m.group(2).strip()

    return f"__auto_{auto_index}__", text


def group_field_photos(field_photos: list) -> list[dict]:
    groups: list[dict] = []
    index: dict[str, int] = {}
    auto_counter = 1

    for raw in field_photos or []:
        path = raw.get("path") or raw.get("file")
        comment = raw.get("comment", "")
        section, caption = parse_photo_comment(comment, auto_counter)
        if section.startswith("__auto_"):
            section = f"현장사진 ({auto_counter})"
            auto_counter += 1

        if section not in index:
            index[section] = len(groups)
            groups.append({"section": section, "items": []})
        groups[index[section]]["items"].append({"path": path, "caption": caption})

    return groups


# ---------------------------------------------------------------------------
# HWP 호환 docx 보고서 생성
# ---------------------------------------------------------------------------

def clean_markdown_for_hwp(text: str) -> str:
    if not text:
        return ""
    cleaned = []
    for line in text.split("\n"):
        line = line.strip().replace("**", "")
        line = re.sub(r"^#{1,6}\s*", "", line)
        line = re.sub(r"^\*\s+", "- ", line)
        if line.startswith("## "):
            line = line.replace("## ", "■ ")
        elif line.startswith("### "):
            line = line.replace("### ", "  ○ ")
        elif line.startswith("* "):
            line = line.replace("* ", "  - ")
        cleaned.append(line)
    return "\n".join(cleaned)


def set_korean_font(run, size: int | None = None, bold: bool | None = None, color: str | None = None) -> None:
    run.font.name = "Malgun Gothic"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_paragraph_spacing(paragraph, before: int = 0, after: int = 0, line: float = 1.15) -> None:
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 100, start: int = 120, bottom: int = 100, end: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_cm: list[float]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(int(Cm(width).emu / 635) for width in widths_cm)))

    grid = table._tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths_cm:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(int(Cm(width).emu / 635)))
        grid.append(grid_col)

    for row in table.rows:
        for idx, width in enumerate(widths_cm):
            cell = row.cells[idx]
            cell.width = Cm(width)
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(int(Cm(width).emu / 635)))


def _add_caption_cell(cell, caption: str) -> None:
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_margins(cell, top=90, bottom=90)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, after=0, line=1.05)
    run = p.add_run(caption or " ")
    set_korean_font(run, size=9, color="3F4F63")


def _add_photo_to_cell(cell, img_path: str | None, width_cm: float) -> None:
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_margins(cell, top=110, bottom=110, start=110, end=110)
    if img_path and os.path.exists(img_path):
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(p, after=0)
        run = p.add_run()
        run.add_picture(img_path, width=Cm(width_cm))
    else:
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("[사진 없음]")
        set_korean_font(run, size=10, color="6B7280")


def add_section_heading(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=8, after=5)
    run = p.add_run(text)
    set_korean_font(run, size=12, bold=True, color="153B5C")


def add_body_text(doc: Document, text: str) -> None:
    for raw in text.splitlines():
        line = raw.strip()
        if not line:
            continue
        p = doc.add_paragraph()
        set_paragraph_spacing(p, after=3, line=1.25)
        if re.match(r"^\d+\.", line):
            p.paragraph_format.left_indent = Cm(0.35)
            p.paragraph_format.first_line_indent = Cm(-0.35)
        run = p.add_run(line)
        set_korean_font(run, size=10, color="27384A")


def create_fieldwork_report(report_data: dict, output_filename: str) -> str:
    location_name = str(report_data.get("location_name") or "사하구 관내")
    task_category = str(report_data.get("task_category") or "현장 점검")
    main_comment = str(report_data.get("main_comment") or "")
    map_image = report_data.get("map_image_path")
    photo_groups = report_data.get("photo_groups") or []
    ai_refined = report_data.get("ai_refined_content") or ""

    doc = Document()
    section = doc.sections[0]
    section.start_type = WD_SECTION_START.NEW_PAGE
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.6)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.7)
    section.right_margin = Cm(1.7)

    style = doc.styles["Normal"]
    style.font.name = "Malgun Gothic"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    style.font.size = Pt(10)

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(title_p, after=4)
    title_run = title_p.add_run(
        f"위치도 및 현장사진({location_name}, {task_category})"
    )
    set_korean_font(title_run, size=15, bold=True, color="111827")

    meta_p = doc.add_paragraph()
    meta_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_paragraph_spacing(meta_p, after=8)
    meta_run = meta_p.add_run(f"작성일자: {date.today().strftime('%Y-%m-%d')}")
    set_korean_font(meta_run, size=9, color="6B7280")

    if map_image and os.path.exists(map_image):
        add_section_heading(doc, "위치도")
        map_p = doc.add_paragraph()
        map_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(map_p, after=8)
        map_p.add_run().add_picture(map_image, width=Cm(16.2))

    for group in photo_groups:
        section_title = group["section"]
        items = group["items"]

        add_section_heading(doc, section_title)

        for i in range(0, len(items), 2):
            chunk = items[i : i + 2]
            cols = len(chunk)
            table = doc.add_table(rows=2, cols=cols)
            table.style = "Table Grid"
            widths = [8.0, 8.0] if cols == 2 else [16.0]
            set_table_geometry(table, widths)

            for col_idx, item in enumerate(chunk):
                img_cell = table.cell(0, col_idx)
                set_cell_shading(img_cell, "F8FAFC")
                _add_photo_to_cell(
                    img_cell,
                    item.get("sanitized_path") or item.get("path"),
                    width_cm=7.35 if cols == 2 else 15.25,
                )
                cap_cell = table.cell(1, col_idx)
                cap_cell.text = ""
                set_cell_shading(cap_cell, "F3F6FA")
                _add_caption_cell(cap_cell, item.get("caption") or "")

            spacer = doc.add_paragraph()
            set_paragraph_spacing(spacer, after=4)

    if main_comment:
        add_section_heading(doc, "주요 의견")
        mc_p = doc.add_paragraph()
        mc_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(mc_p, after=8, line=1.2)
        mc_run = mc_p.add_run(main_comment)
        set_korean_font(mc_run, size=11, bold=True, color="111827")

    if ai_refined:
        add_section_heading(doc, "AI 현장 분석 의견")
        add_body_text(doc, clean_markdown_for_hwp(ai_refined))

    closing_p = doc.add_paragraph()
    closing_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(closing_p, before=10)
    closing_run = closing_p.add_run("위와 같이 사하구 시설물 현장 점검 결과를 보고합니다.")
    set_korean_font(closing_run, size=10, color="27384A")

    out_path = Path(output_filename)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    return str(out_path.resolve())


# ---------------------------------------------------------------------------
# Gemini 분석 Core
# ---------------------------------------------------------------------------

def run_gemini_analysis(
    image_path: str | None,
    location_name: str,
    field_memo: str,
    main_comment: str,
    photo_captions: list[str],
) -> str:
    if not client:
        return "GEMINI_API_KEY가 설정되지 않았습니다. ai/.env 파일을 확인하세요."

    if not image_path or not os.path.exists(image_path):
        return "분석할 현장 사진이 없습니다."

    captions_text = "\n".join(f"- {c}" for c in photo_captions if c) or "(없음)"
    prompt = f"""
당신은 부산광역시 사하구청의 현장 조사 전문 공무원입니다.
'{location_name}' 현장 사진과 담당자 메모를 바탕으로
공식 행정 보고서에 넣을 '현장 상황 요약 및 조치 의견'을 작성하세요.

[작성 원칙]
1. 공문서체 (~함, ~임, ~요망)만 사용할 것.
2. 사진에서 보이는 상세 피해 상태 및 위험도를 행정 공무원 관점으로 기술할 것.
3. 기입된 현장 메모와 메인 코멘트 요구사항을 공문서 단어로 각색하여 누락 없이 반영할 것.

[현장 메모]: {field_memo or "(없음)"}
[메인 코멘트]: {main_comment or "(없음)"}
[사진별 설명]:
{captions_text}

[공식 보고서 내용]:
"""
    try:
        image_object = Image.open(image_path)
        response = client.models.generate_content(
            model="gemini-2.5-flash",
            contents=[image_object, prompt],
        )
        return (response.text or "").strip()
    except Exception as e:
        return f"GenAI API 오류: {e}"


# ---------------------------------------------------------------------------
# 파이프라인 가동
# ---------------------------------------------------------------------------

def safe_filename(name: str) -> str:
    return re.sub(r'[\\/*?:"<>|]', "_", name).strip() or "report"


def run_pipeline(payload: dict) -> dict:
    location_name = payload.get("location_name") or "사하구"
    task_category = payload.get("task_category") or "현장 점검"
    field_memo = payload.get("field_memo") or ""
    main_comment = payload.get("main_comment") or ""

    output_dir = payload.get("output_dir") or "output"
    if not Path(output_dir).is_absolute():
        output_dir = AI_DIR / output_dir
    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)

    work_dir = output_dir / "_work"
    work_dir.mkdir(parents=True, exist_ok=True)

    map_sanitized = sanitize_image(
        payload.get("map_image"),
        str(work_dir / "map.jpg"),
        canvas_size=(1920, 1080),
        fit_mode="cover",
    )

    sanitized_photos = []
    captions = []
    for idx, raw in enumerate(payload.get("field_photos") or []):
        src = raw.get("path") or raw.get("file")
        temp = str(work_dir / f"field_{idx}.jpg")
        sanitized = sanitize_image(src, temp, canvas_size=(1600, 1100), fit_mode="cover")
        _, caption = parse_photo_comment(raw.get("comment", ""), idx + 1)
        captions.append(caption)
        sanitized_photos.append(
            {
                "path": resolve_path(src),
                "sanitized_path": sanitized,
                "comment": raw.get("comment", ""),
                "caption": caption,
            }
        )

    groups = []
    section_index: dict[str, int] = {}
    auto_counter = 1
    for p in sanitized_photos:
        section, caption = parse_photo_comment(p["comment"], auto_counter)
        if section.startswith("__auto_"):
            section = f"현장사진 ({auto_counter})"
            auto_counter += 1
        if section not in section_index:
            section_index[section] = len(groups)
            groups.append({"section": section, "items": []})
        groups[section_index[section]]["items"].append(
            {
                "sanitized_path": p["sanitized_path"],
                "caption": caption,
            }
        )

    primary_image = next(
        (p["sanitized_path"] for p in sanitized_photos if p["sanitized_path"]),
        map_sanitized,
    )

    print("[Engine] Gemini 분석 시작...", file=sys.stderr)
    ai_refined_content = run_gemini_analysis(
        primary_image,
        location_name,
        field_memo,
        main_comment,
        captions,
    )

    report_name = (
        f"위치도_및_현장사진_{safe_filename(location_name)}_{date.today():%Y%m%d}.hwp"
    )
    report_path = create_fieldwork_report(
        {
            "location_name": location_name,
            "task_category": task_category,
            "main_comment": main_comment,
            "map_image_path": map_sanitized,
            "photo_groups": groups,
            "ai_refined_content": ai_refined_content,
        },
        str(output_dir / report_name),
    )

    for temp in work_dir.glob("*.jpg"):
        try:
            temp.unlink()
        except OSError:
            pass

    return {
        "ok": True,
        "location_name": location_name,
        "task_category": task_category,
        "main_comment": main_comment,
        "field_memo": field_memo,
        "ai_refined_content": ai_refined_content,
        "report_file": report_path,
        "error": None,
    }


def emit_result(result: dict) -> None:
    line = RESULT_PREFIX + json.dumps(result, ensure_ascii=False)
    print(line)


def run_pipeline_safe(payload: dict) -> dict:
    try:
        return run_pipeline(payload)
    except Exception as e:
        return {
            "ok": False,
            "ai_refined_content": None,
            "report_file": None,
            "error": str(e),
        }


# ---------------------------------------------------------------------------
# CLI 진입
# ---------------------------------------------------------------------------

if __name__ == "__main__":
    print("=" * 57, file=sys.stderr)
    print(" Saha-gu AI Report Engine", file=sys.stderr)
    print("=" * 57, file=sys.stderr)

    payload = load_payload()
    result = run_pipeline_safe(payload)

    if result.get("ok"):
        print(f"\n[완료] 보고서: {result.get('report_file')}", file=sys.stderr)
        print(f"[AI 본문]\n{result.get('ai_refined_content')}", file=sys.stderr)
    else:
        print(f"\n[실패] {result.get('error')}", file=sys.stderr)

    emit_result(result)
    sys.exit(0 if result.get("ok") else 1)
